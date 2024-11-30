from docx import Document

import requests

import os

# Configuração da API

subscription_key = "zzzz_adicionar_chave_api"

endpoint = 'https://api.cognitive.microsofttranslator.com'

location = "eastus"

language_destination = 'pt-br'

def translator_text(text, target_language):

  # Função de tradução

  headers = {'Ocp-Apim-Subscription-Key': subscription_key, 'Content-type': 'application/json'}

  params = {'api-version': "3.0", 'from': 'en', 'to': target_language}

  body = [{'text': text}]

  response = requests.post(endpoint + "/translate", params=params, headers=headers, json=body)

  return response.json()[0]["translations"][0]["text"]

def translate_document(path):

  # Função para traduzir um documento

  document = Document(path)

  translated_doc = Document()

  for paragraph in document.paragraphs:

    translated_text = translator_text(paragraph.text, language_destination)

    translated_doc.add_paragraph(translated_text)

  translated_doc.save(path.replace(".docx", "_translated.docx"))
